import io
import json
import os
import re
from datetime import datetime
import pdfplumber
from PIL import Image
from models import ExamEntry, ParsedTimetable
from services.datetime_utils import normalize_date, normalize_time, find_date_anchors


ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-opus-4-8")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "nvidia/nemotron-3-super-120b-a12b:free")
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
LLM_TIMEOUT = float(os.getenv("LLM_TIMEOUT_SECONDS", "45"))

_anthropic_client = None
_openrouter_client = None
_gemini_client = None


def get_anthropic_client():
    """Lazily construct the Anthropic client so a missing API key surfaces as a
    clean request-time error, not a startup crash."""
    global _anthropic_client
    if _anthropic_client is None:
        if not os.getenv("ANTHROPIC_API_KEY"):
            raise ValueError(
                "ANTHROPIC_API_KEY is not set. Add it to backend/.env to parse timetables."
            )
        import anthropic

        _anthropic_client = anthropic.Anthropic()
    return _anthropic_client


def get_openrouter_client():
    """Lazily construct the OpenRouter client (OpenAI-compatible) so a missing
    API key surfaces as a clean request-time error, not a startup crash."""
    global _openrouter_client
    if _openrouter_client is None:
        if not os.getenv("OPENROUTER_API_KEY"):
            raise ValueError(
                "OPENROUTER_API_KEY is not set. Add it to backend/.env to parse timetables."
            )
        from openai import OpenAI

        _openrouter_client = OpenAI(
            base_url=OPENROUTER_BASE_URL,
            api_key=os.getenv("OPENROUTER_API_KEY"),
        )
    return _openrouter_client


def get_gemini_client():
    """Lazily construct the Gemini client used as a fallback when OpenRouter fails."""
    global _gemini_client
    if _gemini_client is None:
        if not os.getenv("GEMINI_API_KEY"):
            raise ValueError("GEMINI_API_KEY is not set.")
        from google import genai

        # google-genai auto-reads GOOGLE_API_KEY, not GEMINI_API_KEY — pass it.
        _gemini_client = genai.Client(
            api_key=os.getenv("GEMINI_API_KEY"),
            http_options={"timeout": int(LLM_TIMEOUT * 1000)},  # ms
        )
    return _gemini_client


OPENROUTER_HEADERS = {
    "HTTP-Referer": os.getenv("OPENROUTER_SITE_URL", "http://localhost:5173"),
    "X-Title": os.getenv("OPENROUTER_SITE_NAME", "ExamSync"),
}


def _call_openrouter(prompt: str) -> str:
    """Single-shot JSON extraction. Forces JSON output, deterministic sampling,
    and excludes reasoning tokens (not needed for a one-turn parse). Retries
    without the strict params if the model rejects them."""
    client = get_openrouter_client()
    messages = [{"role": "user", "content": prompt}]
    try:
        completion = client.chat.completions.create(
            model=OPENROUTER_MODEL,
            messages=messages,
            max_tokens=8192,
            temperature=0,
            response_format={"type": "json_object"},
            extra_headers=OPENROUTER_HEADERS,
            extra_body={"reasoning": {"exclude": True}},
            timeout=LLM_TIMEOUT,
        )
    except Exception:
        # Some models reject response_format / reasoning — retry plainly.
        completion = client.chat.completions.create(
            model=OPENROUTER_MODEL,
            messages=messages,
            max_tokens=8192,
            temperature=0,
            extra_headers=OPENROUTER_HEADERS,
            timeout=LLM_TIMEOUT,
        )
    # OpenRouter can return HTTP 200 with an error body and no choices
    # (model unavailable, rate-limited, etc.) — surface that instead of crashing.
    if not getattr(completion, "choices", None):
        err = getattr(completion, "error", None) or getattr(completion, "model_extra", None)
        raise RuntimeError(f"OpenRouter returned no choices: {err}")
    return (completion.choices[0].message.content or "").strip()


# Structured-output schema: the API guarantees the response is valid JSON
# matching this shape, so no code-fence/<think> stripping is ever needed.
EXAM_JSON_SCHEMA = {
    "type": "object",
    "properties": {
        "exams": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "course_code": {"type": "string"},
                    "course_name": {"anyOf": [{"type": "string"}, {"type": "null"}]},
                    "date": {"type": "string"},
                    "time": {"type": "string"},
                    "duration_minutes": {"anyOf": [{"type": "integer"}, {"type": "null"}]},
                    "venue": {"anyOf": [{"type": "string"}, {"type": "null"}]},
                },
                "required": [
                    "course_code", "course_name", "date",
                    "time", "duration_minutes", "venue",
                ],
                "additionalProperties": False,
            },
        }
    },
    "required": ["exams"],
    "additionalProperties": False,
}


def _call_claude(prompt: str) -> str:
    client = get_anthropic_client()
    response = client.with_options(timeout=LLM_TIMEOUT).messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=8192,
        thinking={"type": "adaptive"},
        output_config={"format": {"type": "json_schema", "schema": EXAM_JSON_SCHEMA}},
        messages=[{"role": "user", "content": prompt}],
    )
    return next((b.text for b in response.content if b.type == "text"), "").strip()


def _call_gemini(prompt: str) -> str:
    response = get_gemini_client().models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
    )
    return (response.text or "").strip()


EXTRACTION_PROMPT = """You are an exam timetable parser for students. Timetables
come in MANY layouts and from MANY countries. Adapt to whatever you are given.

Common layouts (handle ALL of them):
- GRID/MATRIX: rows are dates, columns are time slots from a header row; a cell
  holds one or more course code + title pairs. The exam's DATE is its row; its
  TIME is the column header above it.
- LIST / one-row-per-exam: each line/row has a course, a date and a time
  together (e.g. "CIT216 | Programming | 01/06/2026 | 09:00 | Hall A").
- SECTIONED: a date heading followed by the exams under it until the next date.

Date formats vary by country — interpret correctly, then ALWAYS output ISO:
- "01/06/2026" is usually D/M/Y (1 June 2026); "06/01/2026" in US style is
  1 June too only if the day > 12 disambiguates — use surrounding dates and
  weekday names to decide the order, then be consistent across the whole table.
- Textual months in any order: "1 June 2026", "June 1, 2026", "1 Jun 26".
- 2-digit years: "26" -> 2026.
- If the year is missing, infer it from nearby dates or context.

{courses_section}

Return a single JSON object: {{"exams": [ ... ]}}
Each exam object must have exactly these fields:
- course_code: string, exactly as written (e.g. "CIT216")
- course_name: string or null
- date: string in YYYY-MM-DD format (ISO 8601), e.g. "2026-06-01"
- time: string in HH:MM 24h format (8:30am -> "08:30", 11am -> "11:00", 2pm -> "14:00")
- duration_minutes: integer or null, default 120 if not specified
- venue: string or null

Rules:
- Output EVERY exam you can identify; never invent codes or dates.
- Only output the JSON object, no explanation.

Timetable text:
{text}
"""


def build_prompt(raw_text: str, registered: list[str]) -> str:
    if registered:
        codes = ", ".join(registered)
        courses_section = (
            "ONLY return exams for these registered courses (match the course code "
            "ignoring spaces, punctuation and case). Ignore every other course:\n"
            f"{codes}"
        )
    else:
        courses_section = "Extract EVERY exam in the timetable. Do not omit any course."
    return EXTRACTION_PROMPT.format(text=raw_text, courses_section=courses_section)


def _ocr_pdf_pages(pdf) -> str:
    """OCR a scanned/image PDF by rasterizing each page. Best-effort: needs the
    optional tesseract binary; returns '' if OCR isn't available."""
    try:
        import pytesseract  # noqa: F401
    except Exception:
        return ""
    parts = []
    for page in pdf.pages:
        try:
            im = page.to_image(resolution=200).original
            import pytesseract as _pt
            parts.append(_pt.image_to_string(im))
        except Exception as e:
            print(f"[parse] OCR failed on a page: {e}")
    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # Try extracting tables first for structured data
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        text_parts.append(" | ".join(str(c) for c in row if c))
            # Also grab raw text
            raw = page.extract_text()
            if raw:
                text_parts.append(raw)

        text = "\n".join(text_parts)

        # Scanned/image-only PDFs yield little or no extractable text. Fall back
        # to OCR; if that isn't available, raise a clear, actionable error.
        if len(text.strip()) < 40:
            ocr_text = _ocr_pdf_pages(pdf)
            if len(ocr_text.strip()) >= 40:
                return ocr_text
            raise ValueError(
                "This PDF has no selectable text — it looks like a scanned image. "
                "Try exporting a text-based PDF, or upload a clear screenshot/photo "
                "(PNG/JPG) of the timetable instead."
            )

    return text


def extract_text_from_image(file_bytes: bytes) -> str:
    try:
        import pytesseract
    except ImportError:
        raise ValueError(
            "Image OCR is not available in this deployment. "
            "Please upload a PDF, Excel, CSV, or text file instead."
        )
    image = Image.open(io.BytesIO(file_bytes))
    return pytesseract.image_to_string(image)


def extract_text_from_excel(file_bytes: bytes) -> str:
    import pandas as pd
    sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    return "\n\n".join(
        f"Sheet: {name}\n{df.to_string(index=False)}" for name, df in sheets.items()
    )


def extract_text_from_csv(file_bytes: bytes) -> str:
    import pandas as pd
    df = pd.read_csv(io.BytesIO(file_bytes))
    return df.to_string(index=False)


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))
    return "\n".join(text_parts)


def extract_text_from_plaintext(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def extract_raw_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("png", "jpg", "jpeg", "webp", "tiff", "bmp"):
        return extract_text_from_image(file_bytes)
    elif ext in ("xlsx", "xls"):
        return extract_text_from_excel(file_bytes)
    elif ext == "csv":
        return extract_text_from_csv(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    elif ext in ("txt", "text"):
        return extract_text_from_plaintext(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def normalize_course_lines(raw_text: str) -> list[str]:
    courses = []
    seen = set()
    for line in re.split(r"[\n,;]+", raw_text or ""):
        cleaned = re.sub(r"\s+", " ", line).strip(" -\t")
        if not cleaned:
            continue
        key = cleaned.casefold()
        if key not in seen:
            seen.add(key)
            courses.append(cleaned)
    return courses


def parse_json_object(content: str) -> dict:
    content = content.strip()
    # Reasoning models (e.g. nemotron) may leak a <think>...</think> block into
    # the content before the JSON. Strip it defensively.
    content = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
    if content.startswith("```"):
        parts = content.split("```")
        content = parts[1] if len(parts) > 1 else content
        if content.lstrip().startswith("json"):
            content = content.lstrip()[4:]

    try:
        parsed = json.loads(content)
    except json.JSONDecodeError:
        start = content.find("{")
        end = content.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise
        parsed = json.loads(content[start : end + 1])

    if isinstance(parsed, list):
        return {"exams": parsed, "registered_courses": [], "unmatched_courses": []}
    if not isinstance(parsed, dict):
        raise ValueError("Model returned JSON, but not an object.")
    return parsed


def _norm_code(value: str) -> str:
    """Strip everything but alphanumerics and uppercase, so 'CSC 201', 'csc-201'
    and 'CSC201' all compare equal."""
    return re.sub(r"[^a-z0-9]", "", (value or "").lower())


def match_exams(
    exams: list[ExamEntry], registered: list[str]
) -> tuple[list[ExamEntry], list[str]]:
    """Filter exams to those matching a registered course, and report which
    registered courses had no exam. Matching is done on normalized course codes
    and names, tolerant of spacing/punctuation/case."""
    if not registered:
        return exams, []

    matched: list[ExamEntry] = []
    matched_regs: set[str] = set()

    for exam in exams:
        code = _norm_code(exam.course_code)
        name = _norm_code(exam.course_name or "")
        for reg in registered:
            r = _norm_code(reg)
            if not r:
                continue
            if r == code or r in code or code in r or (name and r in name):
                matched.append(exam)
                matched_regs.add(reg)
                break

    unmatched = [reg for reg in registered if reg not in matched_regs]
    return matched, unmatched


# --- Deterministic (no-AI) fallback extractor -------------------------------

_MONTHS = {
    m.lower(): i
    for i, m in enumerate(
        ["January", "February", "March", "April", "May", "June", "July",
         "August", "September", "October", "November", "December"],
        start=1,
    )
}

_DATE_RE = re.compile(
    r"(\d{1,2})(?:st|nd|rd|th)?\s+"
    r"(January|February|March|April|May|June|July|August|September|October|November|December)"
    r"\.?,?\s*(\d{4})?",
    re.IGNORECASE,
)
_CODE_RE = re.compile(r"\b([A-Za-z]{2,4})\s?(\d{3,4})\b")
_TIME_TOKEN_RE = re.compile(r"\b(\d{1,2}(?::\d{2})?\s*(?:am|pm))\b", re.IGNORECASE)


def _detect_time_columns(text: str) -> list[str]:
    """Find the ordered exam times from the timetable header (e.g. 8:30am, 11am,
    2pm -> ['08:30', '11:00', '14:00'])."""
    for line in text.splitlines():
        tokens = _TIME_TOKEN_RE.findall(line)
        if len(tokens) >= 2:
            return [normalize_time(t.replace(" ", "")) for t in tokens]
    out: list[str] = []
    for t in _TIME_TOKEN_RE.findall(text[:2000]):
        nt = normalize_time(t.replace(" ", ""))
        if nt and nt not in out:
            out.append(nt)
    return out


def manual_extract(raw_text: str, registered: list[str]) -> list[ExamEntry]:
    """No-AI fallback. Locates each registered course code in the timetable grid
    and infers its date (nearest preceding date row) and time (column position).
    Best-effort — anything imperfect can be fixed in the editable review table."""
    if not registered:
        return []

    # Repair a year that wrapped onto the next line: "June,\n2026" -> "June 2026".
    text = re.sub(
        r"(January|February|March|April|May|June|July|August|September|October|November|December)"
        r"\.?,?\s*\n\s*(\d{4})",
        r"\1 \2",
        raw_text,
        flags=re.IGNORECASE,
    )

    times = _detect_time_columns(text)
    reg_norm = {_norm_code(r) for r in registered if _norm_code(r)}

    exams: list[ExamEntry] = []
    found: set[str] = set()
    current_date: str | None = None
    current_year: int | None = None

    for line in text.splitlines():
        dm = _DATE_RE.search(line)
        if dm:
            day = int(dm.group(1))
            month = _MONTHS[dm.group(2).lower()]
            if dm.group(3):
                current_year = int(dm.group(3))
            if current_year:
                try:
                    current_date = f"{current_year:04d}-{month:02d}-{day:02d}"
                except Exception:
                    pass

        # Collect genuine course codes, filtering false positives like a month +
        # year ("June 2026") so column indexing stays aligned.
        code_cells = []
        for m in _CODE_RE.finditer(line):
            letters, digits = m.group(1), m.group(2)
            if letters.lower() in _MONTHS:
                continue
            if len(digits) == 4 and digits[:2] in ("19", "20"):  # looks like a year
                continue
            code_cells.append(f"{letters}{digits}")

        # Each code maps to a time column by its order (0 -> 1st time, …).
        for pair_index, code_raw in enumerate(code_cells):
            cn = _norm_code(code_raw)
            if cn in reg_norm and cn not in found:
                if pair_index < len(times) and times[pair_index]:
                    time = times[pair_index]
                elif times:
                    time = times[0]
                else:
                    time = "09:00"
                exams.append(
                    ExamEntry(
                        course_code=code_raw.upper(),
                        course_name=None,
                        date=current_date or "",
                        time=time,
                        duration_minutes=120,
                        venue=None,
                    )
                )
                found.add(cn)

    return exams


def focus_timetable_text(raw_text: str, registered: list[str], max_chars: int = 16000) -> str:
    """Shrink the timetable to only the rows the AI needs: the time-header row,
    every date row (for context), and any row containing a registered course.
    Cuts a 55k-char document to a few KB, so the model responds far faster.
    Falls back to the full text if filtering leaves too little."""
    if not registered:
        return raw_text

    reg = {_norm_code(r) for r in registered if _norm_code(r)}
    kept: list[str] = []
    for line in raw_text.splitlines():
        keep = bool(_TIME_TOKEN_RE.search(line) or _DATE_RE.search(line))
        if not keep:
            for m in _CODE_RE.finditer(line):
                if _norm_code(f"{m.group(1)}{m.group(2)}") in reg:
                    keep = True
                    break
        if keep:
            kept.append(line)

    focused = "\n".join(kept)
    if len(focused) < 40:  # filtered too aggressively — keep the original
        return raw_text
    return focused[:max_chars]


def _exams_from_content(content: str) -> list[ExamEntry]:
    data = parse_json_object(content)
    exams = []
    for entry in data.get("exams", []):
        if not isinstance(entry, dict):
            continue
        # Models sometimes emit placeholder/empty rows (e.g. {"": ""}); skip any
        # row missing the required fields rather than failing the whole parse.
        if not entry.get("course_code") or not entry.get("date") or not entry.get("time"):
            continue
        try:
            exams.append(ExamEntry(**entry))
        except Exception:
            continue
    return exams


def parse_with_claude(raw_text: str, registered: list[str]) -> tuple[list[ExamEntry], str]:
    """Extract exams from the timetable text. Tries Claude first (schema-guaranteed
    JSON via structured outputs); if it errors OR returns zero exams, falls back
    to Gemini, then OpenRouter. Returns (exams, model_used)."""
    prompt = build_prompt(raw_text, registered)
    providers = [
        ("claude", _call_claude, ANTHROPIC_MODEL),
        ("gemini", _call_gemini, GEMINI_MODEL),
        ("openrouter", _call_openrouter, OPENROUTER_MODEL),
    ]

    errors = []
    last_model = ANTHROPIC_MODEL
    for name, call, model in providers:
        try:
            content = call(prompt)
        except Exception as e:
            errors.append(f"{name}: {e}")
            continue
        last_model = model
        try:
            exams = _exams_from_content(content)
        except Exception as e:
            errors.append(f"{name} (bad JSON): {e}")
            continue
        if exams:
            return exams, model
        # Got a valid response but zero exams — try the next provider.

    if len(errors) == len(providers):
        raise RuntimeError("All parsers failed. " + " | ".join(errors))
    return [], last_model


def build_timetable_date_index(raw_text: str) -> dict[str, list[str]]:
    """Map every course code in the timetable to the date of the block it
    appears under. This is the ground truth we check the AI against.

    Works by character offset rather than line by line: we locate every date
    (textual in any order, or numeric/ISO — see datetime_utils.find_date_anchors)
    and every course code, then assign each code to the nearest preceding date.
    This survives the inconsistent line wrapping PDF extraction produces and
    handles timetables from any locale's date format."""
    text = raw_text or ""

    # 1. All date anchors as (offset, ISO-date), locale-agnostic.
    date_points = find_date_anchors(text)
    if not date_points:
        return {}

    # 2. For each course code, find the closest date anchor at or before it.
    index: dict[str, list[str]] = {}
    offsets = [p[0] for p in date_points]
    import bisect

    for m in _CODE_RE.finditer(text):
        letters, digits = m.group(1), m.group(2)
        if letters.lower() in _MONTHS:
            continue
        if len(digits) == 4 and digits[:2] in ("19", "20"):  # looks like a year
            continue
        pos = m.start()
        i = bisect.bisect_right(offsets, pos) - 1
        if i < 0:
            continue  # code appears before the first date — header noise
        iso = date_points[i][1]
        code = _norm_code(f"{letters}{digits}")
        # Record only the FIRST occurrence's date. PDF extraction emits each day's
        # cells in document order first (the authoritative grid row); later raw-text
        # duplicates of the same code would otherwise add spurious dates.
        if code not in index:
            index[code] = [iso]

    return index


def verify_and_correct_dates(
    exams: list[ExamEntry], raw_text: str
) -> tuple[list[ExamEntry], list[str]]:
    """Confirm each extracted exam's date against the uploaded timetable.

    For every exam we look up its course code in the timetable's own date index:
      - date matches the timetable           -> mark verified, no change
      - timetable shows a single, different  -> CORRECT the date to the timetable
        date                                    and record a warning
      - code appears under several dates     -> leave as-is, flag for review
      - code not found in the timetable text -> leave as-is, flag as unverified

    Returns the (possibly corrected) exams and a list of human-readable warnings.
    """
    warnings: list[str] = []
    if not exams:
        return exams, warnings

    index = build_timetable_date_index(raw_text)

    # The verifier only helps for grid/row layouts where codes sit under a date
    # block. For list-style or unusual timetables the index can't anchor the
    # codes — if it covers too few of the extracted exams we skip verification
    # entirely rather than drown the user in false "unverified" flags. The AI
    # result still stands; we just don't second-guess it.
    covered = sum(1 for e in exams if _norm_code(e.course_code) in index)
    if not index or covered / len(exams) < 0.5:
        return exams, warnings

    for exam in exams:
        code = _norm_code(exam.course_code)
        tt_dates = index.get(code)

        if not tt_dates:
            exam.date_verified = False
            exam.date_note = "Course code not found in the timetable — date not verified."
            warnings.append(
                f"{exam.course_code}: not found in the timetable text; please confirm its date."
            )
            continue

        if exam.date in tt_dates:
            exam.date_verified = True
            exam.date_note = None
            continue

        if len(tt_dates) == 1:
            old = exam.date
            exam.date = tt_dates[0]
            exam.date_verified = True
            exam.date_note = f"Date corrected from {old or 'blank'} to match the timetable."
            warnings.append(
                f"{exam.course_code}: date {old or 'blank'} did not match the timetable; "
                f"corrected to {tt_dates[0]}."
            )
        else:
            exam.date_verified = False
            exam.date_note = (
                "Course appears under multiple dates in the timetable "
                f"({', '.join(tt_dates)}); please confirm."
            )
            warnings.append(
                f"{exam.course_code}: found under multiple dates {tt_dates}; please confirm."
            )

    return exams, warnings


def check_duplicate_dates(exams: list[ExamEntry]) -> list[str]:
    """Flag any courses that landed on the same date. Each course should sit on
    its own exam day — two codes sharing a date usually means the extractor
    mis-anchored one of them, so mark both for review instead of letting a
    wrong date reach the calendar silently."""
    warnings: list[str] = []
    by_date: dict[str, list[ExamEntry]] = {}
    for exam in exams:
        if exam.date:
            by_date.setdefault(exam.date, []).append(exam)

    for date, group in by_date.items():
        if len(group) < 2:
            continue
        codes = ", ".join(e.course_code for e in group)
        warnings.append(
            f"{codes} all fall on {date} — confirm each date; courses usually "
            "sit on separate exam days."
        )
        for exam in group:
            exam.date_verified = False
            note = f"Shares {date} with: " + ", ".join(
                e.course_code for e in group if e is not exam
            )
            exam.date_note = f"{exam.date_note} {note}".strip() if exam.date_note else note

    return warnings


def normalize_entries(exams: list[ExamEntry]) -> list[ExamEntry]:
    """Clean recoverable date/time values; leave unrecoverable ones as-is so the
    user can fix them in the editable review table."""
    for exam in exams:
        nd = normalize_date(exam.date)
        nt = normalize_time(exam.time)
        if nd:
            exam.date = nd
        if nt:
            exam.time = nt
    return exams


def parse_timetable(
    filename: str,
    file_bytes: bytes,
    registered_courses_text: str = "",
    registered_courses_filename: str | None = None,
    registered_courses_bytes: bytes | None = None,
) -> ParsedTimetable:
    raw_text = extract_raw_text(filename, file_bytes)
    course_text_parts = [registered_courses_text.strip()] if registered_courses_text.strip() else []
    if registered_courses_filename and registered_courses_bytes:
        course_text_parts.append(extract_raw_text(registered_courses_filename, registered_courses_bytes))
    course_text = "\n".join(part for part in course_text_parts if part)
    registered = normalize_course_lines(course_text) if course_text else []

    exams: list[ExamEntry] = []
    model_used = None

    # Fast path: the instant deterministic extractor. If it cleanly resolves
    # EVERY registered course (with a date and time), skip the AI entirely.
    if registered:
        manual = manual_extract(raw_text, registered)
        if (
            len(manual) == len(registered)
            and all(e.date and e.time for e in manual)
        ):
            exams, model_used = manual, "manual-extractor"

    # Otherwise use AI — on a focused subset of the timetable so it's fast.
    ai_error: str | None = None
    if not exams:
        try:
            exams, model_used = parse_with_claude(
                focus_timetable_text(raw_text, registered), registered
            )
        except Exception as e:
            print(f"[parse] AI extraction failed, using manual fallback: {e}")
            exams, model_used = [], None
            ai_error = str(e)
        # Last resort: partial manual result rather than nothing.
        if not exams and registered:
            manual = manual_extract(raw_text, registered)
            if manual:
                exams, model_used = manual, "manual-extractor"

    # Nothing came back. If the AI providers themselves failed (bad/missing API
    # keys, rate limits, model unavailable), say so clearly instead of returning
    # an empty timetable the user can't diagnose.
    if not exams and ai_error:
        lowered = ai_error.lower()
        if any(s in lowered for s in ("api key", "authentication", "401", "invalid_argument", "unauthorized")):
            raise ValueError(
                "AI service rejected the request — check that OPENROUTER_API_KEY and "
                "GEMINI_API_KEY in backend/.env are valid (not the placeholder values). "
                f"Details: {ai_error}"
            )
        raise ValueError(f"Could not extract any exams from this file. Details: {ai_error}")

    exams = normalize_entries(exams)
    matched, unmatched = match_exams(exams, registered)

    # Cross-check every extracted date against the uploaded timetable and correct
    # any the AI got wrong, so a bad date never silently reaches the calendar.
    matched, date_warnings = verify_and_correct_dates(matched, raw_text)
    date_warnings += check_duplicate_dates(matched)

    print(
        f"[parse] {filename} | model={model_used} | matched {len(matched)}/{len(registered)} "
        f"| date_warnings={len(date_warnings)}"
    )

    return ParsedTimetable(
        exams=matched,
        registered_courses=registered,
        unmatched_courses=unmatched,
        raw_text=raw_text,
        model_used=model_used,
        date_warnings=date_warnings,
    )
